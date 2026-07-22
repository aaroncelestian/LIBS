"""
LIBS identification report as a Word (.docx) document.

For each ranked element: tabulated values plus spectra of the two most intense
matched lines.

Example (CLI):
  .venv/bin/python publication_report.py path/to/spectrum.txt -o reports/out.docx
"""

from __future__ import annotations

import argparse
import io
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from matplotlib.figure import Figure

from identify_elements import (
    ElementHit,
    Match,
    Spectrum,
    find_spectrum_peaks,
    load_line_library,
    load_spectrum,
    match_peaks,
    score_elements,
)
from matplotlib_config import apply_matplotlib_config

ROOT = Path(__file__).resolve().parent
DEFAULT_LIBRARY = ROOT / "nist_lines" / "libs_line_library.csv"

# Fixed report content: two strongest matched lines per element
REPORT_N_LINES = 2


def strongest_matches(hit: ElementHit, n: int = 5) -> list[Match]:
    """Return up to n matches ranked by observed peak intensity (then prominence)."""
    return sorted(
        hit.matches,
        key=lambda m: (m.peak.intensity, m.peak.prominence),
        reverse=True,
    )[:n]


def _window_slice(
    wavelength_nm: np.ndarray,
    intensity: np.ndarray,
    center_nm: float,
    half_width_nm: float,
) -> tuple[np.ndarray, np.ndarray]:
    lo = center_nm - half_width_nm
    hi = center_nm + half_width_nm
    mask = (wavelength_nm >= lo) & (wavelength_nm <= hi)
    return wavelength_nm[mask], intensity[mask]


def plot_element_line_panels(
    spectrum: Spectrum,
    hit: ElementHit,
    *,
    n_lines: int = 5,
    half_width_nm: float = 1.5,
    atmosphere: str | None = None,
    title_extra: str = "",
    fig: Figure | None = None,
) -> Figure:
    """
    One figure for an element: up to ``n_lines`` panels, each zoomed on a strong match.
    Unused panel slots are omitted (figure width scales with panel count).

    Pass an existing ``fig`` to redraw into an embedded GUI canvas.
    """
    apply_matplotlib_config()
    matches = strongest_matches(hit, n=n_lines)
    n = len(matches)

    if fig is None:
        if n == 0:
            fig = plt.figure(figsize=(8, 3))
        else:
            fig_w = max(3.2 * n, 6.0)
            fig = plt.figure(figsize=(fig_w, 3.6))
    else:
        fig.clear()

    if n == 0:
        fig.text(0.5, 0.5, f"{hit.element}: no matched lines", ha="center", va="center")
        fig.tight_layout()
        return fig

    axes = fig.subplots(1, n, sharey=False)
    if n == 1:
        axes = [axes]

    sample = spectrum.meta.path.stem
    atm = f" [{atmosphere}]" if atmosphere else ""
    fig.suptitle(
        f"{sample}{atm} — {hit.element}  "
        f"top {n} intense matched lines  "
        f"({hit.n_peaks} total, confidence {hit.confidence:.0f}%)"
        f"{title_extra}",
        fontsize=11,
        y=0.98,
    )

    wl_all = spectrum.wavelength_nm
    y_all = spectrum.intensity

    for ax, m, i in zip(axes, matches, range(1, n + 1)):
        center = m.peak.wavelength_nm
        x, y = _window_slice(wl_all, y_all, center, half_width_nm)
        if len(x) == 0:
            ax.set_title(f"({i}) no data near {center:.2f} nm", fontsize=9)
            continue

        ax.plot(x, y, color="#1a1a1a", lw=0.9)
        ax.axvline(m.peak.wavelength_nm, color="#c0392b", lw=1.1, alpha=0.85, label="observed")
        ax.axvline(m.line.wavelength_nm, color="#1e8449", lw=1.0, ls="--", alpha=0.9, label="NIST")
        ax.scatter(
            [m.peak.wavelength_nm],
            [m.peak.intensity],
            s=36,
            c="#c0392b",
            zorder=3,
            edgecolors="white",
            linewidths=0.5,
        )

        nist_i = (
            f", NIST I={m.line.intensity:.0f}"
            if m.line.intensity is not None
            else ""
        )
        ax.set_title(
            f"({i}) {m.line.species}  λ={m.peak.wavelength_nm:.3f} nm\n"
            f"I={m.peak.intensity:.0f}  Δλ={m.delta_nm:+.3f} nm{nist_i}",
            fontsize=9,
        )
        ax.set_xlabel("Wavelength (nm)")
        ax.set_xlim(center - half_width_nm, center + half_width_nm)
        y_hi = float(np.max(y)) if len(y) else m.peak.intensity
        y_lo = float(np.min(y)) if len(y) else 0.0
        pad = 0.08 * (y_hi - y_lo + 1.0)
        ax.set_ylim(y_lo - pad, y_hi + pad)
        ax.tick_params(labelsize=8)

    axes[0].set_ylabel("Intensity (counts)")
    axes[0].legend(loc="upper right", fontsize=7, framealpha=0.9)
    fig.tight_layout()
    return fig


def _fig_to_png_bytes(fig: Figure, *, dpi: int = 160) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _line_value_cell(m: Match | None) -> str:
    if m is None:
        return "—"
    nist = (
        f" NIST I={m.line.intensity:.0f}"
        if m.line.intensity is not None
        else ""
    )
    return (
        f"{m.line.species} {m.peak.wavelength_nm:.3f} nm\n"
        f"I={m.peak.intensity:.0f}  Δλ={m.delta_nm:+.3f} nm{nist}"
    )


def export_element_report_docx(
    spectrum: Spectrum,
    hits: list[ElementHit],
    out_path: Path,
    *,
    half_width_nm: float = 1.5,
    atmosphere: str | None = None,
    max_elements: int | None = None,
) -> Path:
    """
    Write a Word document: element value table + two strongest-line spectra each.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for Word reports. "
            "Install with: pip install python-docx"
        ) from exc

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    selected = hits if max_elements is None else hits[:max_elements]
    if not selected:
        raise ValueError("No elements to export — run Find peaks + match first.")

    apply_matplotlib_config()
    doc = Document()

    title = doc.add_heading("LIBS element identification report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    m = spectrum.meta
    atm = atmosphere or "—"
    meta_lines = [
        f"Sample: {m.path.name}",
        f"Config: {m.cfg_path.name if m.cfg_path else '—'}",
        f"Atmosphere: {atm}",
    ]
    if m.laser_energy_mJ is not None:
        meta_lines.append(f"Laser: {m.laser_energy_mJ:g} mJ")
    if m.qs_delay_us is not None:
        meta_lines.append(f"QS delay: {m.qs_delay_us:g} µs")
    if m.integration_time_us is not None:
        delay = (
            f", delay {m.integration_delay_us:g} µs"
            if m.integration_delay_us is not None
            else ""
        )
        meta_lines.append(f"Gate: {m.integration_time_us:g} µs{delay}")
    if m.n_accumulations is not None:
        meta_lines.append(f"Accumulations: {m.n_accumulations}")
    wl0 = float(spectrum.wavelength_nm.min())
    wl1 = float(spectrum.wavelength_nm.max())
    meta_lines.append(f"Range: {wl0:.2f}–{wl1:.2f} nm")

    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    doc.add_heading("Element values", level=1)
    note = doc.add_paragraph(
        "Confidence is identification strength (not concentration). "
        "Line 1 / Line 2 are the two most intense matched peaks."
    )
    note.runs[0].italic = True

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Element"
    hdr[1].text = "Peaks"
    hdr[2].text = "Conf %"
    hdr[3].text = "Line 1 (most intense)"
    hdr[4].text = "Line 2"

    for hit in selected:
        top = strongest_matches(hit, n=REPORT_N_LINES)
        row = table.add_row().cells
        row[0].text = hit.element
        row[1].text = str(hit.n_peaks)
        row[2].text = f"{hit.confidence:.0f}"
        row[3].text = _line_value_cell(top[0] if len(top) > 0 else None)
        row[4].text = _line_value_cell(top[1] if len(top) > 1 else None)

    doc.add_paragraph()
    doc.add_heading("Spectra — two strongest lines per element", level=1)
    legend = doc.add_paragraph(
        "Solid red = observed peak; dashed green = NIST wavelength."
    )
    legend.runs[0].italic = True

    for hit in selected:
        top = strongest_matches(hit, n=REPORT_N_LINES)
        doc.add_heading(
            f"{hit.element}  ·  {hit.n_peaks} peaks  ·  {hit.confidence:.0f}% conf",
            level=2,
        )
        vals = doc.add_paragraph()
        if top:
            bits = [
                f"{m.line.species} {m.peak.wavelength_nm:.3f} nm (I={m.peak.intensity:.0f})"
                for m in top
            ]
            vals.add_run(" · ".join(bits))
        else:
            vals.add_run("No matched lines.")

        fig = plot_element_line_panels(
            spectrum,
            hit,
            n_lines=REPORT_N_LINES,
            half_width_nm=half_width_nm,
            atmosphere=atmosphere,
        )
        png = _fig_to_png_bytes(fig)
        doc.add_picture(io.BytesIO(png), width=Inches(6.5))

    doc.save(out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export LIBS identification report as a Word (.docx) file"
    )
    parser.add_argument("spectrum", type=Path)
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output DOCX path (default: reports/<stem>_lines.docx)",
    )
    parser.add_argument("--library", type=Path, default=DEFAULT_LIBRARY)
    parser.add_argument("--tol-nm", type=float, default=0.12)
    parser.add_argument("--prominence", type=float, default=0.015)
    parser.add_argument("--half-width", type=float, default=1.5)
    parser.add_argument("--atmosphere", type=str, default=None)
    parser.add_argument("--max-elements", type=int, default=None)
    args = parser.parse_args()

    spectrum = load_spectrum(args.spectrum)
    library = load_line_library(args.library)
    peaks = find_spectrum_peaks(spectrum, min_prominence_frac=args.prominence)
    matches = match_peaks(peaks, library, tol_nm=args.tol_nm)
    hits = score_elements(matches, min_peaks=2)

    out = args.output
    if out is None:
        out = ROOT / "reports" / f"{spectrum.meta.path.stem}_lines.docx"

    path = export_element_report_docx(
        spectrum,
        hits,
        out,
        half_width_nm=args.half_width,
        atmosphere=args.atmosphere,
        max_elements=args.max_elements,
    )
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
